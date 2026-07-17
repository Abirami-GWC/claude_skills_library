#!/usr/bin/env python3
"""
generate_brief_docx.py

Optional helper: renders a completed research brief (already drafted as
structured text/markdown by Claude) into a simple .docx file for sharing with a
sales team. Only run this when the user explicitly wants a downloadable file --
for a normal chat answer, just deliver the markdown-formatted brief inline.

This script intentionally does NOT do any research itself -- it only formats
content that has already been gathered and structured. Requires python-docx
(pip install python-docx --break-system-packages).

Usage:
  python3 generate_brief_docx.py brief.json output.docx

Input format (brief.json):
{
  "company_name": "Acme Corp",
  "date": "2026-07-16",
  "company_snapshot": [{"label": "What they do", "text": "...", "status": "Verified", "source": "..."}],
  "signals": [{"title": "...", "detail": "...", "status": "Verified", "source": "...", "date": "..."}],
  "contact_snapshot": [{"label": "Name / Role", "text": "..."}],
  "pain_points": [{"text": "...", "reasoning": "..."}],
  "relevance_to_offer": "...",
  "gaps": ["..."],
  "next_step": "..."
}
"""
import json
import sys

try:
    from docx import Document
except ImportError:
    print("Missing dependency. Install with: pip install python-docx --break-system-packages")
    sys.exit(1)


def build_doc(data):
    doc = Document()
    doc.add_heading(f"Prospect Research Brief — {data.get('company_name', '')}", level=1)
    doc.add_paragraph(f"Compiled: {data.get('date', '')}")

    doc.add_heading("Company Snapshot", level=2)
    for item in data.get("company_snapshot", []):
        p = doc.add_paragraph()
        p.add_run(f"{item['label']}: ").bold = True
        p.add_run(f"{item['text']} ({item.get('status', '')}"
                   f"{', source: ' + item['source'] if item.get('source') else ''})")

    doc.add_heading("Recent Signals", level=2)
    for sig in data.get("signals", []):
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{sig['title']}: ").bold = True
        p.add_run(f"{sig['detail']} ({sig.get('status', '')}"
                   f"{', ' + sig['date'] if sig.get('date') else ''})")

    if data.get("contact_snapshot"):
        doc.add_heading("Contact Snapshot", level=2)
        for item in data["contact_snapshot"]:
            p = doc.add_paragraph()
            p.add_run(f"{item['label']}: ").bold = True
            p.add_run(item["text"])

    if data.get("pain_points"):
        doc.add_heading("Likely Pain Points (inferred)", level=2)
        for pp in data["pain_points"]:
            doc.add_paragraph(f"{pp['text']} — {pp.get('reasoning', '')}", style="List Bullet")

    if data.get("relevance_to_offer"):
        doc.add_heading("Relevance to Offer", level=2)
        doc.add_paragraph(data["relevance_to_offer"])

    if data.get("gaps"):
        doc.add_heading("Gaps / Could Not Verify", level=2)
        for g in data["gaps"]:
            doc.add_paragraph(g, style="List Bullet")

    if data.get("next_step"):
        doc.add_heading("Suggested Next Step", level=2)
        doc.add_paragraph(data["next_step"])

    return doc


def main():
    if len(sys.argv) != 3:
        print("Usage: python3 generate_brief_docx.py brief.json output.docx")
        sys.exit(1)

    data = json.loads(open(sys.argv[1]).read())
    doc = build_doc(data)
    doc.save(sys.argv[2])
    print(f"Saved brief to {sys.argv[2]}")


if __name__ == "__main__":
    main()
