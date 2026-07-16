#!/usr/bin/env python3
"""
generate_battlecard_docx.py

Renders a populated competitor-profile JSON (matching
knowledge/competitor_profile_schema.json) into a compact, one-page Word (.docx)
battlecard, following the section order in knowledge/battlecard_structure.md.

Usage:
    python generate_battlecard_docx.py <path-to-profile.json> <output.docx>

Dependencies: python-docx (preinstalled in this environment).

After generating, verify it actually fits one page — convert to PDF and check
page count:
    soffice --headless --convert-to pdf output.docx
    pdfinfo output.pdf | grep Pages
If it overflows to a second page, trim content (see Best Practices in
SKILL.md) rather than re-running with a smaller font size.
"""

import json
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


NAVY = RGBColor(0x1A, 0x2B, 0x4C)
GRAY = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xB0, 0x2A, 0x2A)


def add_heading(doc, text, size=13, color=NAVY, space_before=8, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.size = Pt(9)
        p.add_run(" — " + text).font.size = Pt(9)
    else:
        p.add_run(text).font.size = Pt(9)
    return p


def build_docx(profile: dict, out_path: str):
    doc = Document()

    # Compact margins for one-page density
    section = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(0.5))

    # Header
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run(f"Battlecard: Us vs. {profile.get('competitor_name', '[Competitor]')}")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = NAVY

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"As of {profile.get('as_of_date', '[date]')}  |  Scope: {profile.get('segment_scope', 'general')}"
    )
    meta_run.italic = True
    meta_run.font.size = Pt(8)
    meta_run.font.color.rgb = GRAY

    # Snapshot table
    add_heading(doc, "Snapshot")
    snap = profile.get("snapshot", {})
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, key in [
        ("Founded", "founded"), ("HQ", "hq"), ("Size/Funding", "size_or_funding"),
        ("Target customer", "target_customer_profile"), ("Primary use case", "primary_use_case"),
        ("Deployment", "deployment_model"),
    ]:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(snap.get(key, "") or "")
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    # Strengths
    add_heading(doc, "Their Strengths (real)")
    for s in profile.get("genuine_strengths", []):
        add_bullet(doc, s.get("claim", ""))

    # Weaknesses
    add_heading(doc, "Their Weaknesses / Landmines")
    for w in profile.get("known_weaknesses", []):
        claim = w.get("claim", "")
        q = w.get("landmine_question", "")
        text = claim + (f'  (Ask: "{q}")' if q else "")
        add_bullet(doc, text)

    # Head-to-head
    add_heading(doc, "Head-to-Head Snapshot")
    h2h = profile.get("head_to_head", [])
    if h2h:
        t2 = doc.add_table(rows=1, cols=3)
        t2.style = "Light Grid Accent 1"
        hdr = t2.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Dimension", "Us", "Them"
        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(8)
        for row_data in h2h:
            row = t2.add_row().cells
            row[0].text = row_data.get("dimension", "")
            row[1].text = row_data.get("us", "")
            row[2].text = row_data.get("them", "")
            for cell in row:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)

    # Differentiators
    add_heading(doc, "Our Key Differentiators")
    for d in profile.get("our_differentiators", []):
        add_bullet(doc, d.get("proof_point", ""), bold_lead=d.get("claim", ""))

    # Objections
    add_heading(doc, "Objection Handling")
    for o in profile.get("objections", []):
        add_bullet(doc, o.get("response", ""), bold_lead=f"\u201c{o.get('objection', '')}\u201d")

    # Talking points
    add_heading(doc, "Talking Points")
    for tp in profile.get("talking_points", []):
        add_bullet(doc, tp)

    # Do not say
    do_not_say = profile.get("do_not_say", [])
    if do_not_say:
        add_heading(doc, "\u26a0 Do NOT Say", color=RED)
        for item in do_not_say:
            add_bullet(doc, item)

    # Footer / sources
    doc.add_paragraph().add_run("")
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        f"Pricing confidence: {profile.get('pricing', {}).get('confidence', 'unverified')}. "
        f"Refresh this card periodically — competitive facts go stale."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(7)
    footer_run.font.color.rgb = GRAY

    doc.save(out_path)


def main():
    if len(sys.argv) != 3:
        print("Usage: python generate_battlecard_docx.py <path-to-profile.json> <output.docx>")
        sys.exit(2)

    in_path, out_path = sys.argv[1], sys.argv[2]
    with open(in_path, "r", encoding="utf-8") as f:
        profile = json.load(f)

    build_docx(profile, out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
