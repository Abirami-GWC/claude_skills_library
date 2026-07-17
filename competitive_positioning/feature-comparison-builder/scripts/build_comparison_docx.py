#!/usr/bin/env python3
"""
build_comparison_docx.py

Renders a populated comparison-matrix JSON file (see
knowledge/comparison_matrix_schema.json, supports N vendors) into a
presentation-ready Word (.docx) document: one table per dimension category,
with a legend and sourcing footer. Uses landscape orientation automatically
when there are 3+ vendors, since wide tables overflow portrait width.

Usage:
    python build_comparison_docx.py <path-to-matrix.json> <output.docx>

Dependencies: python-docx (preinstalled).

After generating, verify readability for wide (3+ vendor) tables by converting
to PDF and inspecting a rendered page:
    soffice --headless --convert-to pdf output.docx
    pdftoppm -jpeg -r 100 output.pdf page
If columns are still cramped, consider splitting into multiple documents by
category rather than shrinking font size below ~8pt.
"""

import json
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

SYMBOL = {
    "Full": "\u2705",
    "Partial": "\u25D0",
    "Add-on": "\u2795",
    "Unsupported": "\u274C",
    "Unknown": "\u2753",
}

NAVY = RGBColor(0x1A, 0x2B, 0x4C)
GRAY = RGBColor(0x55, 0x55, 0x55)


def set_landscape(doc, wide=False):
    section = doc.sections[0]
    if wide:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(0.6))


def build_docx(matrix: dict, out_path: str):
    doc = Document()

    vendors = matrix.get("vendors", [])
    dimensions = matrix.get("dimensions", [])
    wide = len(vendors) >= 3
    set_landscape(doc, wide=wide)

    title = doc.add_paragraph()
    r = title.add_run(matrix.get("title") or "Feature Comparison")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = NAVY

    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Accurate as of {matrix.get('as_of_date', '[date]')}")
    meta_run.italic = True
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = GRAY

    legend = doc.add_paragraph()
    legend_run = legend.add_run(
        "Legend:  \u2705 Full    \u25D0 Partial    \u2795 Add-on    "
        "\u274C Unsupported    \u2753 Unknown / not verified"
    )
    legend_run.font.size = Pt(8)
    legend_run.font.color.rgb = GRAY

    # Group dimensions by category, one table per category
    categories = []
    for d in dimensions:
        cat = d.get("category", "General")
        if cat not in categories:
            categories.append(cat)

    sources = []

    for cat in categories:
        cat_dims = [d for d in dimensions if d.get("category", "General") == cat]

        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(4)
        hr = heading.add_run(cat)
        hr.bold = True
        hr.font.size = Pt(12)
        hr.font.color.rgb = NAVY

        n_cols = 1 + len(vendors)
        table = doc.add_table(rows=1, cols=n_cols)
        table.style = "Light Grid Accent 1"

        hdr = table.rows[0].cells
        hdr[0].text = "Dimension"
        for i, v in enumerate(vendors, start=1):
            hdr[i].text = v.get("display_name", v.get("id", ""))
        for cell in hdr:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)

        for dim in cat_dims:
            row = table.add_row().cells
            row[0].text = dim.get("label", "")
            cells_by_vendor = {c.get("vendor_id"): c for c in dim.get("cells", [])}
            for i, v in enumerate(vendors, start=1):
                cell_data = cells_by_vendor.get(v.get("id"), {})
                level = cell_data.get("support_level", "Unknown")
                detail = cell_data.get("detail", "")
                symbol = SYMBOL.get(level, "?")
                text = symbol + (f" {detail}" if detail else "")
                row[i].text = text
                if cell_data.get("source_note") or cell_data.get("source_url"):
                    sources.append(
                        f"{v.get('display_name', '')} — {dim.get('label', '')}: "
                        f"{cell_data.get('source_note', '')} {cell_data.get('source_url', '')}".strip()
                    )
            for cell in row:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)

    # Sources appendix
    if sources:
        doc.add_paragraph().add_run("")
        src_heading = doc.add_paragraph()
        sh = src_heading.add_run("Sources")
        sh.bold = True
        sh.font.size = Pt(10)
        sh.font.color.rgb = NAVY
        for s in sources:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(s).font.size = Pt(7)

    doc.save(out_path)


def main():
    if len(sys.argv) != 3:
        print("Usage: python build_comparison_docx.py <path-to-matrix.json> <output.docx>")
        sys.exit(2)

    in_path, out_path = sys.argv[1], sys.argv[2]
    with open(in_path, "r", encoding="utf-8") as f:
        matrix = json.load(f)

    build_docx(matrix, out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
