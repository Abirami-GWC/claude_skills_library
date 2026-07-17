#!/usr/bin/env python3
"""
build_campaign_brief_docx.py

Renders a populated campaign-brief JSON file (matching
knowledge/campaign_brief_schema.json) into a shareable Word (.docx) document:
core idea, audience/goal, messaging pillars, a channel-by-channel table
(angle/format/hook/CTA/metric/owner), excluded channels with reasons, and
open questions — so a channel owner or stakeholder can review or execute
directly from the document.

Usage:
    python build_campaign_brief_docx.py <path-to-brief.json> <output.docx>

Dependencies: python-docx (preinstalled).

This script does not validate strategic quality (e.g. whether channels
actually fit the goal) — that judgment happens in
instructions/validation_rules.md before this script is run. This script only
renders whatever populated data it is given.
"""

import json
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor

NAVY = RGBColor(0x1A, 0x2B, 0x4C)
GRAY = RGBColor(0x55, 0x55, 0x55)


def add_heading(doc, text, size=13, space_before=14, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = NAVY
    return p


def build_docx(brief: dict, out_path: str):
    doc = Document()
    section = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(0.7))

    title = doc.add_paragraph()
    r = title.add_run(brief.get("campaign_name") or "Campaign Brief")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = NAVY

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"Goal: {brief.get('primary_goal', '[goal]')}   |   "
        f"Audience: {brief.get('audience_persona', '[persona]')}   |   "
        f"Timeframe: {brief.get('timeframe', '[timeframe]')}"
    )
    meta_run.italic = True
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = GRAY

    # Core idea
    add_heading(doc, "Core Idea / Asset")
    core = brief.get("core_idea_or_asset", {})
    p = doc.add_paragraph()
    p.add_run(core.get("summary", "")).font.size = Pt(10)
    if core.get("source_link_or_reference"):
        ref = doc.add_paragraph()
        ref_run = ref.add_run(f"Reference: {core['source_link_or_reference']}")
        ref_run.italic = True
        ref_run.font.size = Pt(8)
        ref_run.font.color.rgb = GRAY

    # Messaging pillars
    add_heading(doc, "Messaging Pillars")
    for pillar in brief.get("messaging_pillars", []):
        bp = doc.add_paragraph(style="List Bullet")
        r = bp.add_run(pillar.get("pillar", ""))
        r.bold = True
        r.font.size = Pt(10)
        if pillar.get("supporting_point"):
            bp.add_run("  —  " + pillar["supporting_point"]).font.size = Pt(10)

    # Channel table
    add_heading(doc, "Channel Plan")
    channels = brief.get("channels", [])
    if channels:
        cols = ["Channel", "Angle", "Format", "Hook", "CTA", "Metric", "Owner"]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, c in enumerate(cols):
            hdr[i].text = c
            for p in hdr[i].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(8)
        for ch in channels:
            row = table.add_row().cells
            values = [
                ch.get("channel", ""), ch.get("angle", ""), ch.get("format", ""),
                ch.get("example_hook", ""), ch.get("cta", ""),
                ch.get("success_metric", ""), ch.get("owner", ""),
            ]
            for i, v in enumerate(values):
                row[i].text = v
                for p in row[i].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)

    # Excluded channels
    excluded = brief.get("excluded_channels", [])
    if excluded:
        add_heading(doc, "Channels Considered and Excluded")
        for ex in excluded:
            bp = doc.add_paragraph(style="List Bullet")
            r = bp.add_run(ex.get("channel", ""))
            r.bold = True
            r.font.size = Pt(9)
            bp.add_run("  —  " + ex.get("reason_excluded", "")).font.size = Pt(9)

    # Open questions
    open_qs = brief.get("open_questions", [])
    if open_qs:
        add_heading(doc, "Open Questions")
        for q in open_qs:
            bp = doc.add_paragraph(style="List Bullet")
            bp.add_run(q).font.size = Pt(9)

    doc.save(out_path)


def main():
    if len(sys.argv) != 3:
        print("Usage: python build_campaign_brief_docx.py <path-to-brief.json> <output.docx>")
        sys.exit(2)

    in_path, out_path = sys.argv[1], sys.argv[2]
    with open(in_path, "r", encoding="utf-8") as f:
        brief = json.load(f)

    build_docx(brief, out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
